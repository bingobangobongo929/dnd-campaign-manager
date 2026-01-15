#!/usr/bin/env python3
"""
Comprehensive Vault Character Import Script - COMPLETE EDITION
Extracts ALL character data from Word documents with ZERO data loss.

Handles:
- Backstory (prose and phased)
- TLDR / Summary
- NPCs / Relationships
- Companions / Familiars
- Session Notes / Journal
- Character Writings (letters, stories, poems)
- Quotes and Common Phrases
- Rumors (true/false)
- Plot Hooks / Knives
- External Links (D&D Beyond, YouTube, etc.)
- Player Meta (OOC info)
- DM Q&A
- Possessions / Inventory
- Combat Stats
- Party Relations
- Physical Appearance
- Secondary Characters in same document
"""

import os
import sys
import json
import re
import unicodedata
from datetime import datetime
from pathlib import Path

sys.stdout.reconfigure(encoding='utf-8')

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:
    print("Installing python-docx...")
    os.system(f"{sys.executable} -m pip install python-docx")
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError

try:
    import requests
except ImportError:
    print("Installing requests...")
    os.system(f"{sys.executable} -m pip install requests")
    import requests


# =============================================================================
# CONFIGURATION
# =============================================================================

CHARACTERS_DIR = r"C:\Users\edbar\Downloads\Character\Charactere"
API_URL = "http://localhost:3000/api/vault/import"
OUTPUT_FILE = r"C:\Users\edbar\Downloads\Character\vault_characters_import.json"
DEFAULT_GAME_SYSTEM = "D&D 5e"
DEFAULT_PRONOUNS = "she/her"


# =============================================================================
# TEXT UTILITIES
# =============================================================================

def normalize_text(text: str) -> str:
    """Normalize Unicode characters and clean up text."""
    if not text:
        return ""
    text = unicodedata.normalize('NFKC', text)
    text = text.replace('"', '"').replace('"', '"')
    text = text.replace(''', "'").replace(''', "'")
    text = text.replace('–', '-').replace('—', '-')
    # Remove emojis from name but preserve them in content
    return text.strip()


def strip_emojis_from_name(text: str) -> str:
    """Remove emojis from character names."""
    if not text:
        return text
    # Remove common emoji patterns
    emoji_pattern = re.compile(
        "["
        "\U0001F600-\U0001F64F"  # emoticons
        "\U0001F300-\U0001F5FF"  # symbols & pictographs
        "\U0001F680-\U0001F6FF"  # transport & map symbols
        "\U0001F1E0-\U0001F1FF"  # flags
        "\U00002702-\U000027B0"
        "\U000024C2-\U0001F251"
        "]+", flags=re.UNICODE
    )
    return emoji_pattern.sub('', text).strip()


def fix_formatting(text: str) -> str:
    """Apply markdown formatting fixes while preserving paragraph breaks."""
    if not text:
        return text

    # Preserve double newlines as paragraph breaks
    text = re.sub(r'\n{3,}', '\n\n', text)

    return text.strip()


def clean_backstory_text(text: str) -> str:
    """Remove markdown headers from backstory text to create clean prose."""
    if not text:
        return text

    # Remove ## headers but keep the content
    cleaned = re.sub(r'^#{1,3}\s+[A-Za-z\s\'\"]+\n+', '\n', text, flags=re.MULTILINE)

    # Remove **bold** section headers
    cleaned = re.sub(r'^\*\*[A-Za-z\s\'\"]+\*\*\n+', '\n', cleaned, flags=re.MULTILINE)

    # Clean up multiple newlines while preserving paragraph breaks
    cleaned = re.sub(r'\n{3,}', '\n\n', cleaned)

    return cleaned.strip()


def fix_common_typos(text: str) -> str:
    """Fix common typos."""
    if not text:
        return text
    typos = {
        'Rouge': 'Rogue',
        'rouge': 'rogue',
        'Palyer': 'Player',
        'palyer': 'player',
        'Charcter': 'Character',
        'charcter': 'character',
        'recieve': 'receive',
        'seperate': 'separate',
        'occured': 'occurred',
        'definately': 'definitely',
    }
    for typo, correction in typos.items():
        text = text.replace(typo, correction)
    return text


# =============================================================================
# DOCUMENT EXTRACTION
# =============================================================================

def extract_paragraphs(filepath: str) -> list[str]:
    """Extract all paragraphs from a docx file, preserving structure."""
    try:
        doc = Document(filepath)
    except PackageNotFoundError:
        print(f"  Error: Could not open {filepath}")
        return []

    paragraphs = []

    for para in doc.paragraphs:
        text = normalize_text(para.text)
        if text:
            # Check if it's a heading
            if para.style and para.style.name.startswith('Heading'):
                level = para.style.name.replace('Heading ', '')
                try:
                    level_num = int(level)
                    prefix = '#' * level_num
                    paragraphs.append(f"{prefix} {text}")
                except ValueError:
                    paragraphs.append(f"## {text}")
            else:
                paragraphs.append(text)

    # Also get table content
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paragraphs.append(" | ".join(cells))

    return paragraphs


def extract_full_text(paragraphs: list[str]) -> str:
    """Join paragraphs into full text preserving paragraph breaks."""
    return "\n\n".join(paragraphs)


# =============================================================================
# SECTION DETECTION AND EXTRACTION
# =============================================================================

# All known section headers
SECTION_HEADERS = {
    'backstory': ['backstory', 'background', 'history'],
    'tldr': ['tldr', 'tl;dr', 'backstory highlights', 'highlights'],
    'knives': ['knives', 'plot hooks', 'story hooks'],
    'quotes': ['quotes', 'quotations'],
    'npcs': ['npcs', 'important people', 'how she bonds', 'how he bonds', 'crew member relations'],
    'session': ['session', 'personal session notes', 'personal notes'],
    'letters': ['letters to nana', 'letters', 'letter'],
    'campfire': ['campfire stories', 'campfire'],
    'rumors': ['rumors', 'rumours', '2 rumors'],
    'dm_qa': ["dm's question", 'dm questions', 'dm question'],
    'player_info': ['player info', 'player information'],
    'possessions': ['possessions', 'possession', 'inventory'],
    'tips': ['tips', 'gameplay tips', 'combat tips'],
    'extra': ['extra', 'extra information', 'additional'],
    'weakness': ['weakness', 'weaknesses', 'fears'],
    'description': ['description'],
    'early_life': ['early life', 'childhood'],
    'student_life': ['student life', 'the cracks'],
    'adult_life': ['adult life', 'the new and the boring'],
    'good_path': ['the good path', 'the "good path"'],
    'common_phrases': ['common phrases'],
    'new_friend': ['new friend', 'meeting'],
    'what_knows': ['what she knows', 'what he knows'],
    'turning_point': ['turning point'],
    'rebellion': ['rebellion and escape', 'rebellion'],
    'tension': ['tension and diverging paths', 'tension'],
    'upbringing': ['upbringing'],
}


def is_section_header(text: str) -> bool:
    """Check if text looks like a section header."""
    if not text:
        return False

    # Starts with markdown header
    if text.startswith('#'):
        return True

    lower = text.lower().strip().rstrip(':')

    # Check against known headers
    for headers in SECTION_HEADERS.values():
        for h in headers:
            if lower == h or lower.startswith(h + ' ') or lower.endswith(' ' + h):
                return True

    # Check for Session X pattern
    if re.match(r'^session\s*#?\d+', lower):
        return True

    # Short ALL CAPS titles
    if len(text) < 40 and text.isupper():
        return True

    # Capitalized short titles
    if len(text) < 50 and re.match(r'^[A-Z][a-z]+(\s+[A-Za-z]+){0,4}:?\s*$', text):
        return True

    return False


def find_section(paragraphs: list[str], target_headers: list[str], stop_at_any_header: bool = True) -> str:
    """Extract content from a section, stopping at the next section header."""
    content = []
    in_section = False

    for para in paragraphs:
        lower = para.lower().strip().rstrip(':')

        # Check if this is our target section
        is_target = False
        for header in target_headers:
            if lower == header or lower.startswith(header + ' ') or lower.startswith(header + ':'):
                is_target = True
                break
            if para.startswith('#') and header in lower:
                is_target = True
                break

        if is_target:
            in_section = True
            # Get any text after the header on the same line
            for header in target_headers:
                if lower.startswith(header):
                    rest = para[len(header):].strip().lstrip(':').strip()
                    if rest and not rest.startswith('#'):
                        content.append(rest)
                    break
            continue

        if in_section:
            # Check if we should stop
            if stop_at_any_header and is_section_header(para):
                break
            content.append(para)

    return '\n\n'.join(content)


def find_all_sections_of_type(paragraphs: list[str], section_prefix: str) -> list[dict]:
    """Find all sections that start with a prefix (like 'Session 1', 'Session 2')."""
    sections = []
    current_section = None
    current_content = []

    for para in paragraphs:
        # Check for section start
        match = re.match(rf'^{section_prefix}\s*#?(\d+)', para, re.IGNORECASE)
        if match:
            # Save previous section
            if current_section is not None:
                sections.append({
                    'number': current_section,
                    'content': '\n\n'.join(current_content)
                })

            current_section = int(match.group(1))
            # Get rest of header line
            rest = para[match.end():].strip()
            if rest.startswith('-') or rest.startswith(':'):
                rest = rest[1:].strip()
            current_content = [rest] if rest else []
        elif current_section is not None:
            # Check if this is a different section type
            if is_section_header(para) and not para.lower().startswith(section_prefix.lower()):
                sections.append({
                    'number': current_section,
                    'content': '\n\n'.join(current_content)
                })
                current_section = None
                current_content = []
            else:
                current_content.append(para)

    # Don't forget the last section
    if current_section is not None and current_content:
        sections.append({
            'number': current_section,
            'content': '\n\n'.join(current_content)
        })

    return sections


# =============================================================================
# CHARACTER WRITINGS EXTRACTION
# =============================================================================

def extract_letters(paragraphs: list[str]) -> list[dict]:
    """Extract letters (e.g., 'Letters to Nana' from Fleur)."""
    letters = []
    current_letter = None
    current_content = []
    in_letters_section = False

    letter_titles = [
        'The Apology', 'The first kill', 'The Kidnapping', 'New Feelings',
        'Lost', 'Getting Use To It', 'Love', 'Dear Nana', 'Dear friend'
    ]

    for para in paragraphs:
        lower = para.lower().strip()

        # Detect letters section
        if 'letters to nana' in lower or 'letters' in lower:
            in_letters_section = True
            continue

        # Check for letter title
        is_letter_title = False
        for title in letter_titles:
            if para.strip() == title or para.strip().rstrip(':') == title:
                is_letter_title = True
                # Save previous letter
                if current_letter and current_content:
                    letters.append({
                        'title': current_letter,
                        'type': 'letter',
                        'content': '\n\n'.join(current_content),
                        'recipient': 'Nana'
                    })
                current_letter = title
                current_content = []
                break

        if is_letter_title:
            continue

        # Check for "Dear X" pattern
        if para.strip().startswith('Dear ') and len(para.strip()) < 50:
            if current_letter and current_content:
                letters.append({
                    'title': current_letter,
                    'type': 'letter',
                    'content': '\n\n'.join(current_content),
                    'recipient': 'Nana'
                })
            current_letter = para.strip()
            current_content = []
            continue

        # Collect content
        if current_letter:
            # Check for session header (end of letters section)
            if re.match(r'^session\s*#?\d+', lower):
                letters.append({
                    'title': current_letter,
                    'type': 'letter',
                    'content': '\n\n'.join(current_content),
                    'recipient': 'Nana'
                })
                break
            current_content.append(para)

    # Save last letter
    if current_letter and current_content:
        letters.append({
            'title': current_letter,
            'type': 'letter',
            'content': '\n\n'.join(current_content),
            'recipient': 'Nana'
        })

    return letters


def extract_campfire_stories(paragraphs: list[str]) -> list[dict]:
    """Extract campfire stories."""
    stories = []
    current_story = None
    current_content = []
    in_campfire_section = False

    for para in paragraphs:
        lower = para.lower().strip()

        # Detect campfire section
        if 'campfire stories' in lower:
            in_campfire_section = True
            continue

        if in_campfire_section:
            # Story titles are typically short
            if len(para) < 50 and not para.startswith('#') and is_section_header(para):
                # Check if this is a different major section
                if any(h in lower for h in ['letters', 'session', 'quotes']):
                    break

                # Save previous story
                if current_story and current_content:
                    stories.append({
                        'title': current_story,
                        'type': 'story',
                        'content': '\n\n'.join(current_content)
                    })
                current_story = para.strip()
                current_content = []
            elif current_story:
                current_content.append(para)
            elif not current_story and para and len(para) < 50:
                # First story title
                current_story = para.strip()

    # Save last story
    if current_story and current_content:
        stories.append({
            'title': current_story,
            'type': 'story',
            'content': '\n\n'.join(current_content)
        })

    return stories


# =============================================================================
# RUMORS EXTRACTION
# =============================================================================

def extract_rumors(paragraphs: list[str]) -> list[dict]:
    """Extract rumors with true/false flags."""
    rumors = []
    in_rumors_section = False

    for para in paragraphs:
        lower = para.lower().strip()

        if 'rumors' in lower or 'rumours' in lower:
            in_rumors_section = True
            continue

        if in_rumors_section:
            # Check for end of section
            if is_section_header(para) and 'rumors' not in lower:
                break

            # Parse rumor pattern: A) or B) followed by text and (this is true/false)
            match = re.match(r'^[A-Z]\)\s*(.+?)(?:\s*\(this is (true|not true|false)\))?$', para, re.IGNORECASE)
            if match:
                statement = match.group(1).strip()
                is_true_text = match.group(2) if match.group(2) else ''
                is_true = 'true' in is_true_text.lower() and 'not' not in is_true_text.lower()
                rumors.append({
                    'statement': statement,
                    'is_true': is_true
                })

    return rumors


# =============================================================================
# DM Q&A EXTRACTION
# =============================================================================

def extract_dm_qa(paragraphs: list[str]) -> list[dict]:
    """Extract DM question/answer pairs."""
    qa_pairs = []
    current_question = None
    current_answer = []
    in_qa_section = False

    for para in paragraphs:
        lower = para.lower().strip()

        # Detect Q&A section
        if "dm's question" in lower or 'dm question' in lower:
            in_qa_section = True
            continue

        if in_qa_section:
            # Check for end of section
            if is_section_header(para) and 'dm' not in lower and '?' not in para:
                if current_question and current_answer:
                    qa_pairs.append({
                        'question': current_question,
                        'answer': '\n\n'.join(current_answer)
                    })
                break

            # Check for question (ends with ?)
            if para.strip().endswith('?'):
                # Save previous Q&A
                if current_question and current_answer:
                    qa_pairs.append({
                        'question': current_question,
                        'answer': '\n\n'.join(current_answer)
                    })
                current_question = para.strip()
                current_answer = []
            elif current_question:
                current_answer.append(para)

    # Save last Q&A
    if current_question and current_answer:
        qa_pairs.append({
            'question': current_question,
            'answer': '\n\n'.join(current_answer)
        })

    return qa_pairs


# =============================================================================
# PLAYER META EXTRACTION
# =============================================================================

def extract_player_meta(paragraphs: list[str]) -> dict:
    """Extract player OOC information."""
    meta = {}
    in_player_section = False

    for para in paragraphs:
        lower = para.lower().strip()

        if 'player info' in lower:
            in_player_section = True
            continue

        if in_player_section:
            if is_section_header(para) and 'player' not in lower:
                break

            # Parse key: value patterns
            if ':' in para:
                parts = para.split(':', 1)
                key = parts[0].strip().lower()
                value = parts[1].strip()

                if 'discord' in key:
                    meta['player_discord'] = value
                elif 'timezone' in key or 'location' in key:
                    meta['player_timezone'] = value
                elif 'experience' in key:
                    meta['player_experience'] = value
                elif 'name' in key and 'age' in key:
                    # "Name, Age, Gender: Celina, 26, Female"
                    meta['player_name'] = value
                elif 'fun' in key or 'annoy' in key or 'personality' in key:
                    if 'player_preferences' not in meta:
                        meta['player_preferences'] = {}
                    meta['player_preferences'][key] = value

    return meta


# =============================================================================
# PARTY RELATIONS EXTRACTION
# =============================================================================

def extract_party_relations(paragraphs: list[str]) -> list[dict]:
    """Extract crew member / party relations."""
    relations = []
    current_member = None
    current_notes = []
    in_relations_section = False

    for para in paragraphs:
        lower = para.lower().strip()

        if 'crew member relations' in lower or 'party relations' in lower:
            in_relations_section = True
            continue

        if in_relations_section:
            # Check for end
            if is_section_header(para) and 'relations' not in lower and not is_party_member_name(para):
                if current_member and current_notes:
                    relations.append({
                        'name': current_member,
                        'notes': '\n'.join(current_notes)
                    })
                break

            # Check for party member name (short, capitalized)
            if is_party_member_name(para):
                if current_member and current_notes:
                    relations.append({
                        'name': current_member,
                        'notes': '\n'.join(current_notes)
                    })
                current_member = para.strip()
                current_notes = []
            elif current_member:
                current_notes.append(para)

    if current_member and current_notes:
        relations.append({
            'name': current_member,
            'notes': '\n'.join(current_notes)
        })

    return relations


def is_party_member_name(text: str) -> bool:
    """Check if text looks like a party member name."""
    text = text.strip()
    if len(text) > 30:
        return False
    if len(text) < 3:
        return False
    # Capitalized word(s)
    if re.match(r'^[A-Z][a-z]+(\s+[A-Za-z]+)?$', text):
        return True
    return False


# =============================================================================
# COMMON PHRASES EXTRACTION
# =============================================================================

def extract_common_phrases(paragraphs: list[str]) -> list[str]:
    """Extract common phrases / catchphrases."""
    phrases = []
    in_phrases_section = False

    for para in paragraphs:
        lower = para.lower().strip()

        if 'common phrases' in lower:
            in_phrases_section = True
            # Check for inline phrases: "phrase1" - "phrase2" pattern
            matches = re.findall(r'"([^"]+)"', para)
            phrases.extend(matches)
            continue

        if in_phrases_section:
            if is_section_header(para) and 'phrases' not in lower:
                break

            # Extract quoted phrases
            matches = re.findall(r'"([^"]+)"', para)
            phrases.extend(matches)

    return phrases


# =============================================================================
# COMBAT STATS EXTRACTION
# =============================================================================

def extract_combat_stats(full_text: str) -> dict:
    """Extract combat statistics."""
    stats = {}

    # Kill score
    kill_match = re.search(r'kill\s*score[:\s]+(\d+)', full_text, re.IGNORECASE)
    if kill_match:
        stats['kills'] = int(kill_match.group(1))

    # Inspiration
    insp_match = re.search(r'inspiration[:\s]+(\d+)', full_text, re.IGNORECASE)
    if insp_match:
        stats['inspiration'] = int(insp_match.group(1))

    return stats if stats else None


# =============================================================================
# POSSESSIONS EXTRACTION
# =============================================================================

def extract_possessions(paragraphs: list[str]) -> list[dict]:
    """Extract possessions/inventory items."""
    possessions = []
    in_possessions_section = False

    for para in paragraphs:
        lower = para.lower().strip()

        if 'possession' in lower:
            in_possessions_section = True
            continue

        if in_possessions_section:
            if is_section_header(para) and 'possession' not in lower:
                break

            # Parse item patterns
            # "2 Healing Potion - 2d4+2"
            # "474 gold"
            item_match = re.match(r'^(\d+)\s+(.+?)(?:\s+-\s+(.+))?$', para.strip())
            if item_match:
                possessions.append({
                    'quantity': int(item_match.group(1)),
                    'name': item_match.group(2).strip(),
                    'details': item_match.group(3).strip() if item_match.group(3) else None
                })
            elif para.strip() and not is_section_header(para):
                # Simple item
                possessions.append({
                    'name': para.strip(),
                    'quantity': 1
                })

    return possessions


# =============================================================================
# QUOTES EXTRACTION
# =============================================================================

def extract_quotes_section(paragraphs: list[str]) -> list[str]:
    """Extract quotes from a dedicated quotes section."""
    quotes = []
    in_quotes_section = False

    for para in paragraphs:
        lower = para.lower().strip()

        if lower == 'quotes' or lower == 'quotes:':
            in_quotes_section = True
            continue

        if in_quotes_section:
            if is_section_header(para) and 'quotes' not in lower:
                break

            # Check if this is a philosophical quote
            if para.strip() and len(para) > 10:
                quotes.append(para.strip())

    return quotes


def extract_inline_quotes(full_text: str) -> list[str]:
    """Extract quoted dialogue from text."""
    quotes = []

    # Find all quoted text
    found = re.findall(r'"([^"]{10,200})"', full_text)

    for q in found:
        # Filter out URLs, session references, etc.
        if 'http' in q.lower():
            continue
        if 'session' in q.lower():
            continue
        if q.startswith('The ') and len(q) < 30:  # Likely a title
            continue
        quotes.append(q)

    return quotes[:20]  # Limit to 20


# =============================================================================
# RELATIONSHIPS / NPCs EXTRACTION
# =============================================================================

def extract_relationships(paragraphs: list[str], char_name: str) -> list[dict]:
    """Extract NPC/relationship information."""
    relationships = []
    current_npc = None
    current_details = []
    in_npc_section = False

    # Known NPC section markers
    npc_markers = ['how she bonds', 'how he bonds', 'important people', 'npcs']

    for i, para in enumerate(paragraphs):
        lower = para.lower().strip()

        # Detect NPC section start
        if any(marker in lower for marker in npc_markers):
            in_npc_section = True
            continue

        if in_npc_section:
            # Check if this is an NPC name
            if is_npc_header(para, char_name):
                # Save previous
                if current_npc and current_details:
                    relationships.append(build_relationship(current_npc, current_details))

                current_npc = para.strip()
                # Handle "Name - Title" pattern
                if ' - ' in current_npc:
                    parts = current_npc.split(' - ')
                    current_npc = parts[0].strip()
                current_details = []
            elif current_npc:
                if is_section_header(para) and not is_npc_header(para, char_name):
                    # End of NPC section
                    relationships.append(build_relationship(current_npc, current_details))
                    break
                current_details.append(para)

    # Save last NPC
    if current_npc and current_details:
        relationships.append(build_relationship(current_npc, current_details))

    return relationships


def is_npc_header(text: str, char_name: str) -> bool:
    """Check if text looks like an NPC name header."""
    if not text:
        return False

    # Skip if it contains the main character's name
    if char_name:
        char_first_name = char_name.split()[0].lower()
        if char_first_name in text.lower():
            return False

    # Skip common section words
    section_words = [
        'backstory', 'background', 'early', 'adult', 'student', 'quotes',
        'session', 'notes', 'tldr', 'important', 'npcs', 'bonds', 'path',
        'friend', 'knows', 'life', 'the', 'how', 'what', 'new', 'speaking',
        'knives', 'rumors', 'letters', 'campfire'
    ]
    lower = text.lower()
    if any(lower.startswith(w) for w in section_words):
        return False

    # Skip if starts with descriptive words
    description_starts = [
        'dont', "don't", 'has ', 'is ', 'was ', 'very ', 'needs', 'uses ',
        'he ', 'she ', 'they ', 'his ', 'her ', 'their ', 'interesting',
        'talks ', 'dislikes', 'loves ', 'enjoys ', 'hates ', 'favors',
        'got ', 'from ', 'prev', 'was ', 'when ', 'we ', 'i '
    ]
    if any(lower.startswith(w) for w in description_starts):
        return False

    # Should be relatively short
    if len(text) > 50:
        return False

    # Should look like a name (capitalized)
    name_pattern = r'^[A-Z][a-z]+(\s+(de|von|van|la|el|[A-Z][a-z]+))*(\s*[\(\)"\'\-A-Za-z\.]+)?$'
    if re.match(name_pattern, text.strip()):
        if len(text) > 3:
            return True

    return False


def build_relationship(npc_name: str, details: list[str]) -> dict:
    """Build a relationship object from extracted data."""
    description = ' '.join(details)[:1500]
    details_lower = description.lower()

    rel = {
        'related_name': npc_name,
        'description': description,
        'relationship_type': 'other'
    }

    # Determine relationship type
    if any(w in details_lower for w in ['father', 'dad', 'mother', 'mom', 'parent', 'brother', 'sister', 'sibling', 'twin', 'grandmother', 'nana']):
        rel['relationship_type'] = 'family'
        if 'father' in details_lower or 'dad' in details_lower:
            rel['relationship_label'] = 'Father'
        elif 'mother' in details_lower or 'mom' in details_lower:
            rel['relationship_label'] = 'Mother'
        elif 'brother' in details_lower:
            rel['relationship_label'] = 'Brother'
        elif 'sister' in details_lower:
            rel['relationship_label'] = 'Sister'
        elif 'twin' in details_lower:
            rel['relationship_label'] = 'Twin'
        elif 'grandmother' in details_lower or 'nana' in details_lower:
            rel['relationship_label'] = 'Grandmother'
    elif any(w in details_lower for w in ['patron', 'warlock']):
        rel['relationship_type'] = 'patron'
        rel['relationship_label'] = 'Patron'
    elif any(w in details_lower for w in ['mentor', 'teacher', 'master', 'tutor', 'tutoring']):
        rel['relationship_type'] = 'mentor'
        rel['relationship_label'] = 'Mentor'
    elif any(w in details_lower for w in ['friend', 'ally', 'close']):
        rel['relationship_type'] = 'friend'
        rel['relationship_label'] = 'Friend'
    elif any(w in details_lower for w in ['enemy', 'rival', 'nemesis', "don't trust"]):
        rel['relationship_type'] = 'enemy'
        rel['relationship_label'] = 'Enemy'
    elif any(w in details_lower for w in ['love', 'romantic', 'partner', 'spouse', 'husband', 'wife']):
        rel['relationship_type'] = 'romantic'
        rel['relationship_label'] = 'Romantic Interest'
    elif any(w in details_lower for w in ['familiar', 'pet', 'fam.']):
        rel['relationship_type'] = 'companion'
        rel['relationship_label'] = 'Companion/Familiar'
    elif any(w in details_lower for w in ['baron', 'employer', 'boss', 'captain']):
        rel['relationship_type'] = 'employer'
        rel['relationship_label'] = 'Employer'

    return rel


# =============================================================================
# BACKSTORY PHASES EXTRACTION
# =============================================================================

def extract_backstory_phases(paragraphs: list[str]) -> list[dict]:
    """Extract structured backstory phases."""
    phases = []
    phase_headers = [
        ('Early Life', ['early life', 'childhood', 'youth']),
        ('Student Life', ['student life', 'the cracks', 'training']),
        ('The Good Path', ['the good path', 'the "good path"']),
        ('Upbringing', ['upbringing']),
        ('Tension', ['tension and diverging paths', 'tension']),
        ('Rebellion', ['rebellion and escape', 'rebellion']),
        ('New Life', ['new life and inner conflict', 'new life']),
        ('Turning Point', ['turning point']),
        ('New Path', ['new path']),
        ('Adult Life', ['adult life', 'the new and the boring']),
        ('Current', ['current', 'recent events', 'present day']),
    ]

    for title, headers in phase_headers:
        content = find_section(paragraphs, headers)
        if content and len(content) > 50:
            phases.append({
                'title': title,
                'content': fix_formatting(content)
            })

    return phases


# =============================================================================
# COMPANIONS EXTRACTION
# =============================================================================

def extract_companions(paragraphs: list[str], full_text: str) -> list[dict]:
    """Extract companion/pet/familiar information."""
    companions = []

    # Pattern matching for companions
    patterns = [
        r'(?:familiar|pet|mount|companion)[,:\s]+(?:a\s+)?(?:named\s+)?([A-Z][a-z]+)',
        r'([A-Z][a-z]+)(?:\s+the\s+|\s+is\s+(?:her|his)\s+)(?:familiar|pet|mount|companion)',
        r'(?:named|called)\s+([A-Z][a-z]+)[\s,]+(?:her|his)\s+(?:familiar|pet|companion)',
        r'([A-Z][a-z]+)\s+-\s+Fam\.',  # Penny - Fam.
    ]

    for pattern in patterns:
        matches = re.finditer(pattern, full_text, re.IGNORECASE)
        for match in matches:
            name = match.group(1)
            if name and len(name) > 2 and name.lower() not in ['the', 'her', 'his']:
                # Check if already added
                if not any(c['name'].lower() == name.lower() for c in companions):
                    companions.append({
                        'name': name,
                        'type': 'companion',
                        'description': ''
                    })

    return companions


# =============================================================================
# SESSION JOURNAL EXTRACTION
# =============================================================================

def extract_session_journal(paragraphs: list[str]) -> list[dict]:
    """Extract session journal entries."""
    sessions = find_all_sections_of_type(paragraphs, 'Session')

    journal = []
    for session in sessions:
        if session['content'] and len(session['content']) > 10:
            journal.append({
                'session_number': session['number'],
                'title': f"Session {session['number']}",
                'summary': fix_formatting(session['content'])
            })

    return journal


# =============================================================================
# MEDIA LINKS EXTRACTION
# =============================================================================

def extract_media_links(full_text: str) -> dict:
    """Extract URLs for theme music, character sheets, etc."""
    links = {}

    urls = re.findall(r'https?://[^\s<>"\']+[^\s<>"\',.\)]', full_text)

    for url in urls:
        if 'dndbeyond.com' in url:
            links['character_sheet_url'] = url
        elif 'youtube.com' in url or 'youtu.be' in url:
            if 'theme_music_url' not in links:
                links['theme_music_url'] = url
        elif 'spotify.com' in url:
            if 'spotify_playlist' not in links:
                links['spotify_playlist'] = url

    return links


# =============================================================================
# PHYSICAL APPEARANCE EXTRACTION
# =============================================================================

def extract_physical_appearance(paragraphs: list[str], full_text: str) -> dict:
    """Extract physical appearance details."""
    appearance = {}

    patterns = {
        'height': [r'height[:\s]+([^\n|,]+)', r'(\d+[\'"\s]*\d*["\s]*)(?:\s*tall)?'],
        'weight': [r'weight[:\s]+([^\n|,]+)', r'(\d+\s*(?:lbs?|kg|pounds?))'],
        'hair': [r'hair[:\s]+([^\n|,]+)', r'hair\s*(?:color)?[:\s]*([a-zA-Z]+\s*[a-zA-Z]*)'],
        'eyes': [r'eyes?[:\s]+([^\n|,]+)', r'eye\s*(?:color)?[:\s]*([a-zA-Z]+)'],
        'skin': [r'skin[:\s]+([^\n|,]+)', r'complexion[:\s]+([^\n|,]+)'],
        'voice': [r'voice[:\s]+([^\n|,]+)'],
        'age': [r'age[:\s]+([^\n|,]+)', r'(\d+)\s*years?\s*old'],
    }

    for field, field_patterns in patterns.items():
        for pattern in field_patterns:
            match = re.search(pattern, full_text, re.IGNORECASE)
            if match:
                value = match.group(1).strip()
                if value and len(value) < 100:
                    appearance[field] = value
                    break

    # Look for distinguishing marks
    marks_pattern = r'(?:scar|tattoo|birthmark|marking)[s]?[:\s]+([^\n.]+)'
    match = re.search(marks_pattern, full_text, re.IGNORECASE)
    if match:
        appearance['distinguishing_marks'] = match.group(1).strip()

    return appearance


# =============================================================================
# RACE/CLASS DETECTION
# =============================================================================

def detect_race_class(full_text: str) -> tuple:
    """Detect race, class, and subclass from text."""
    race = None
    char_class = None
    subclass = None

    races = [
        'Human', 'Elf', 'Half-Elf', 'Dwarf', 'Halfling', 'Gnome', 'Half-Orc',
        'Tiefling', 'Dragonborn', 'Aasimar', 'Genasi', 'Goliath', 'Tabaxi',
        'Kenku', 'Firbolg', 'Changeling', 'Warforged', 'Goblin', 'Shadar-Kai',
        'Half-Siren', 'Triton', 'Kalashtar', 'Water Genasi', 'Lightfoot Halfling'
    ]

    classes = [
        'Barbarian', 'Bard', 'Cleric', 'Druid', 'Fighter', 'Monk',
        'Paladin', 'Ranger', 'Rogue', 'Sorcerer', 'Warlock', 'Wizard',
        'Artificer', 'Blood Hunter', 'Blood Mage', 'Blood Archer'
    ]

    # Detect race
    for r in sorted(races, key=len, reverse=True):  # Check longer names first
        if re.search(r'\b' + re.escape(r) + r'\b', full_text, re.IGNORECASE):
            race = r
            break

    # Detect class
    for c in sorted(classes, key=len, reverse=True):
        if re.search(r'\b' + re.escape(c) + r'\b', full_text, re.IGNORECASE):
            char_class = c
            break

    return race, char_class, subclass


def detect_game_system(full_text: str) -> str:
    """Detect the game system from text."""
    lower = full_text.lower()

    if 'talabheim' in lower or 'morr' in lower or 'hexen' in lower or 'warhammer' in lower:
        return 'Warhammer Fantasy'
    elif 'mech' in lower and 'pilot' in lower:
        return 'Lancer'
    elif 'spelljammer' in lower or 'astral sea' in lower:
        return 'Spelljammer/D&D 5e'
    elif 'pathfinder' in lower:
        return 'Pathfinder'
    else:
        return 'D&D 5e'


# =============================================================================
# CHARACTER TAGS GENERATION
# =============================================================================

def extract_character_tags(full_text: str, race: str, char_class: str) -> list[str]:
    """Generate character tags based on content."""
    tags = []
    lower = full_text.lower()

    tag_keywords = {
        'blood-magic': ['blood magic', 'blood pact', 'hemomancy', 'blood archer'],
        'pirate': ['pirate', 'ship captain', 'sailor', 'sea'],
        'changeling': ['changeling'],
        'noble': ['noble', 'baron', 'aristocrat', 'royalty', 'princess'],
        'criminal': ['thief', 'criminal', 'smuggler', 'con artist', 'assassin'],
        'scholar': ['scholar', 'academic', 'researcher', 'library'],
        'military': ['soldier', 'military', 'army', 'veteran'],
        'religious': ['temple', 'priest', 'cleric', 'faith'],
        'nature': ['druid', 'forest', 'animals', 'nature'],
        'arcane': ['wizard', 'sorcerer', 'magic', 'hexen'],
        'cursed': ['curse', 'cursed'],
        'tragic': ['tragic', 'loss', 'grief', 'trauma'],
        'mysterious': ['mysterious', 'secret', 'hidden'],
        'assassin': ['assassin', 'killer', 'umbra'],
        'baker': ['baker', 'cook', 'pie', 'baking'],
        'doctor': ['doctor', 'physician', 'healer', 'medicine'],
    }

    for tag, keywords in tag_keywords.items():
        if any(kw in lower for kw in keywords):
            tags.append(tag)

    if race:
        tags.append(race.lower().replace(' ', '-'))
    if char_class:
        tags.append(char_class.lower())

    return list(set(tags))[:15]


# =============================================================================
# GOLD EXTRACTION
# =============================================================================

def extract_gold(full_text: str) -> int:
    """Extract gold amount."""
    match = re.search(r'(\d+)\s*(?:gold|gp)', full_text, re.IGNORECASE)
    if match:
        return int(match.group(1))
    return None


# =============================================================================
# TLDR / PLOT HOOKS EXTRACTION
# =============================================================================

def extract_bullet_points(text: str) -> list[str]:
    """Extract bullet points from text."""
    if not text:
        return []

    bullets = []
    lines = text.split('\n')

    for line in lines:
        line = line.strip()
        # Check for bullet patterns
        if line.startswith(('- ', '• ', '* ', '· ', '→ ')):
            content = line[2:].strip()
            if content and len(content) > 3:
                bullets.append(content)
        # Check for numbered patterns
        elif re.match(r'^\d+[\.\)]\s+', line):
            content = re.sub(r'^\d+[\.\)]\s+', '', line).strip()
            if content and len(content) > 3:
                bullets.append(content)

    # If no bullets found, try splitting on periods for sentence-style TLDRs
    if not bullets and text:
        for line in lines:
            line = line.strip()
            if line and len(line) > 10 and not is_section_header(line):
                bullets.append(line)

    return bullets


# =============================================================================
# SECONDARY CHARACTERS DETECTION
# =============================================================================

def detect_secondary_characters(paragraphs: list[str], full_text: str) -> list[dict]:
    """Detect if document contains multiple character concepts."""
    secondary = []

    # Look for patterns like "Character Name, Age, Race, Background: Name, 25, Human, Noble"
    char_pattern = r'Character\s+(?:Name|Info)[,:\s]+(?:Age[,:\s]+)?(?:Race[,:\s]+)?(?:Background)?[:\s]+([A-Z][a-z]+),?\s*(\d+)?,?\s*([A-Za-z\-\s]+)?'

    matches = re.finditer(char_pattern, full_text, re.IGNORECASE)
    names_found = []

    for match in matches:
        name = match.group(1)
        if name and name not in names_found:
            names_found.append(name)
            secondary.append({
                'name': name,
                'age': match.group(2) if match.group(2) else None,
                'race': match.group(3).strip() if match.group(3) else None
            })

    # Only return if we found more than one character concept
    if len(secondary) > 1:
        return secondary[1:]  # Return all except the first (main character)

    return []


# =============================================================================
# MAIN CHARACTER EXTRACTION
# =============================================================================

def extract_character(filepath: str) -> dict:
    """Extract ALL character data from a document with ZERO data loss."""
    filename = os.path.basename(filepath)
    name = os.path.splitext(filename)[0]
    name = strip_emojis_from_name(normalize_text(name))

    print(f"\nProcessing: {name}")

    paragraphs = extract_paragraphs(filepath)
    if not paragraphs:
        print(f"  Warning: No paragraphs extracted")
        return None

    full_text = extract_full_text(paragraphs)
    full_text = fix_common_typos(full_text)

    # ============ BASIC INFO ============
    game_system = detect_game_system(full_text)
    race, char_class, subclass = detect_race_class(full_text)
    appearance = extract_physical_appearance(paragraphs, full_text)

    # ============ TEXT SECTIONS ============
    backstory = find_section(paragraphs, ['backstory', 'background', 'history'])
    personality = find_section(paragraphs, ['personality', 'traits'])
    goals = find_section(paragraphs, ['goals', 'objectives', 'ambitions'])
    secrets = find_section(paragraphs, ['secrets', 'hidden', 'secret'])
    fears_text = find_section(paragraphs, ['fears', 'phobias', 'weakness'])
    extra = find_section(paragraphs, ['extra', 'extra information', 'additional'])

    # Summary from TLDR or first paragraph
    summary = find_section(paragraphs, ['summary', 'overview'])
    if not summary and backstory:
        first_para = backstory.split('\n\n')[0]
        if len(first_para) > 50:
            summary = first_para[:500] + '...' if len(first_para) > 500 else first_para

    # ============ ARRAYS ============
    tldr_section = find_section(paragraphs, ['tldr', 'tl;dr', 'backstory highlights'])
    tldr = extract_bullet_points(tldr_section)

    knives_section = find_section(paragraphs, ['knives', 'plot hooks', 'story hooks'])
    plot_hooks = extract_bullet_points(knives_section)

    open_questions_section = find_section(paragraphs, ['open questions', 'mysteries'])
    open_questions = extract_bullet_points(open_questions_section)

    weaknesses = extract_bullet_points(fears_text) if fears_text else []
    fears = extract_bullet_points(find_section(paragraphs, ['fears']))

    # ============ QUOTES ============
    quotes_from_section = extract_quotes_section(paragraphs)
    quotes_inline = extract_inline_quotes(full_text)
    all_quotes = list(set(quotes_from_section + quotes_inline))

    common_phrases = extract_common_phrases(paragraphs)

    # ============ STRUCTURED DATA ============
    relationships = extract_relationships(paragraphs, name)
    backstory_phases = extract_backstory_phases(paragraphs)
    companions = extract_companions(paragraphs, full_text)
    session_journal = extract_session_journal(paragraphs)

    # ============ NEW: CHARACTER WRITINGS ============
    letters = extract_letters(paragraphs)
    campfire_stories = extract_campfire_stories(paragraphs)
    character_writings = letters + campfire_stories

    # ============ NEW: RUMORS ============
    rumors = extract_rumors(paragraphs)

    # ============ NEW: DM Q&A ============
    dm_qa = extract_dm_qa(paragraphs)

    # ============ NEW: PLAYER META ============
    player_meta = extract_player_meta(paragraphs)

    # ============ NEW: PARTY RELATIONS ============
    party_relations = extract_party_relations(paragraphs)

    # ============ NEW: POSSESSIONS ============
    possessions = extract_possessions(paragraphs)

    # ============ NEW: COMBAT STATS ============
    combat_stats = extract_combat_stats(full_text)

    # ============ TAGS AND LINKS ============
    tags = extract_character_tags(full_text, race, char_class)
    media_links = extract_media_links(full_text)
    gold = extract_gold(full_text)

    # ============ CLEAN BACKSTORY FOR NOTES FIELD ============
    # The UI displays 'notes' field as "Full Backstory"
    # Use clean backstory text without ## headers
    if backstory:
        notes = clean_backstory_text(fix_formatting(backstory))
    else:
        notes = clean_backstory_text(fix_formatting(full_text))

    # ============ BUILD CHARACTER OBJECT ============
    character = {
        'name': name,
        'type': 'pc',
        'game_system': game_system,
        'pronouns': DEFAULT_PRONOUNS,

        # Basic info
        'race': race,
        'class': char_class,
        'subclass': subclass,

        # Physical appearance
        'height': appearance.get('height'),
        'weight': appearance.get('weight'),
        'hair': appearance.get('hair'),
        'eyes': appearance.get('eyes'),
        'skin': appearance.get('skin'),
        'voice': appearance.get('voice'),
        'age': appearance.get('age'),
        'distinguishing_marks': appearance.get('distinguishing_marks'),

        # Text content
        'backstory': clean_backstory_text(fix_formatting(backstory)) if backstory else None,
        'description': clean_backstory_text(fix_formatting(backstory)) if backstory else None,
        'summary': fix_formatting(summary) if summary else None,
        'personality': fix_formatting(personality) if personality else None,
        'goals': fix_formatting(goals) if goals else None,
        'secrets': fix_formatting(secrets) if secrets else None,
        'notes': fix_formatting(notes) if notes else None,

        # Arrays
        'quotes': all_quotes if all_quotes else None,
        'common_phrases': common_phrases if common_phrases else None,
        'tldr': tldr if tldr else None,
        'plot_hooks': plot_hooks if plot_hooks else None,
        'open_questions': open_questions if open_questions else None,
        'weaknesses': weaknesses if weaknesses else None,
        'fears': fears if fears else None,
        'character_tags': tags if tags else None,
        'gameplay_tips': None,  # Could extract from "Tips" sections

        # Structured JSONB
        'backstory_phases': backstory_phases if backstory_phases else None,
        'companions': companions if companions else None,
        'session_journal': session_journal if session_journal else None,
        'possessions': possessions if possessions else None,

        # NEW: Character writings
        'character_writings': character_writings if character_writings else None,

        # NEW: Rumors
        'rumors': rumors if rumors else None,

        # NEW: DM Q&A
        'dm_qa': dm_qa if dm_qa else None,

        # NEW: Player meta
        'player_discord': player_meta.get('player_discord'),
        'player_timezone': player_meta.get('player_timezone'),
        'player_experience': player_meta.get('player_experience'),
        'player_preferences': player_meta.get('player_preferences'),
        'player_name': player_meta.get('player_name'),

        # NEW: Party relations
        'party_relations': party_relations if party_relations else None,

        # NEW: Combat stats
        'combat_stats': combat_stats,

        # Relationships (for separate table)
        'relationships': relationships if relationships else None,

        # Media links
        'theme_music_url': media_links.get('theme_music_url'),
        'character_sheet_url': media_links.get('character_sheet_url'),
        'spotify_playlist': media_links.get('spotify_playlist'),

        # Tracking
        'gold': gold,
        'status': 'Active' if len(full_text) > 200 else 'Concept',
        'source_file': filename,
        'imported_at': datetime.now().isoformat(),
        'raw_document_text': full_text,
    }

    # ============ REPORT ============
    print(f"  Game System: {game_system}")
    print(f"  Race: {race or 'Unknown'}")
    print(f"  Class: {char_class or 'Unknown'}")
    print(f"  Backstory: {len(backstory or '')} chars")
    print(f"  Relationships: {len(relationships)}")
    print(f"  Backstory Phases: {len(backstory_phases)}")
    print(f"  Session Journal: {len(session_journal)} entries")
    print(f"  Character Writings: {len(character_writings)} (letters/stories)")
    print(f"  Rumors: {len(rumors)}")
    print(f"  DM Q&A: {len(dm_qa)}")
    print(f"  Quotes: {len(all_quotes)}")
    print(f"  Party Relations: {len(party_relations)}")
    print(f"  Tags: {', '.join(tags[:5])}{'...' if len(tags) > 5 else ''}")

    return character


# =============================================================================
# PROCESS DIRECTORY
# =============================================================================

def process_directory(directory: str) -> list[dict]:
    """Process all Word documents in a directory."""
    characters = []

    if not os.path.exists(directory):
        print(f"Error: Directory not found: {directory}")
        return characters

    for filename in sorted(os.listdir(directory)):
        if not filename.endswith('.docx'):
            continue
        if filename.startswith('~'):
            continue
        # Skip the ideas file
        if 'Ideas' in filename or 'ideas' in filename:
            continue

        filepath = os.path.join(directory, filename)
        try:
            char = extract_character(filepath)
            if char:
                characters.append(char)
        except Exception as e:
            print(f"  Error: {e}")
            import traceback
            traceback.print_exc()

    return characters


def send_to_api(characters: list[dict], api_url: str) -> dict:
    """Send characters to the import API."""
    try:
        response = requests.post(
            api_url,
            json={'characters': characters},
            headers={'Content-Type': 'application/json'},
            timeout=120
        )
        return response.json()
    except requests.exceptions.RequestException as e:
        return {'error': str(e)}


# =============================================================================
# MAIN
# =============================================================================

def main():
    print("=" * 70)
    print("Vault Character Import - COMPLETE EDITION")
    print("Zero Data Loss • Full Schema Support")
    print("=" * 70)

    characters = process_directory(CHARACTERS_DIR)

    if not characters:
        print("\nNo characters extracted!")
        return

    print(f"\n{'=' * 70}")
    print(f"Extracted {len(characters)} characters")
    print("=" * 70)

    # Wrap in the format the API expects
    output_data = {
        'characters': characters
    }

    # Save to JSON
    with open(OUTPUT_FILE, 'w', encoding='utf-8') as f:
        json.dump(output_data, f, indent=2, ensure_ascii=False)
    print(f"\nSaved to: {OUTPUT_FILE}")

    # Summary
    print("\n" + "=" * 70)
    print("EXTRACTION SUMMARY")
    print("=" * 70)

    for char in characters:
        print(f"\n{char['name']}:")
        print(f"  Race/Class: {char.get('race', '?')}/{char.get('class', '?')}")
        print(f"  Backstory: {len(char.get('backstory') or '')} chars")
        print(f"  Notes: {len(char.get('notes') or '')} chars")
        print(f"  Relationships: {len(char.get('relationships') or [])}")
        print(f"  Phases: {len(char.get('backstory_phases') or [])}")
        print(f"  Sessions: {len(char.get('session_journal') or [])}")
        print(f"  Writings: {len(char.get('character_writings') or [])}")
        print(f"  Rumors: {len(char.get('rumors') or [])}")
        print(f"  DM Q&A: {len(char.get('dm_qa') or [])}")
        print(f"  Quotes: {len(char.get('quotes') or [])}")

    print("\n" + "=" * 70)
    print("IMPORT INSTRUCTIONS")
    print("=" * 70)
    print(f"\n1. Run the database migration:")
    print("   npx supabase migration up")
    print(f"\n2. Import via the Settings page:")
    print("   http://localhost:3000/settings/import")
    print(f"\n3. Upload this file:")
    print(f"   {OUTPUT_FILE}")


if __name__ == "__main__":
    main()
