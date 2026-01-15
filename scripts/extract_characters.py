#!/usr/bin/env python3
"""
Character Document Extractor
Extracts text from all .docx files and saves as .txt for manual review/import.
"""

import os
import sys
import codecs
import json
from pathlib import Path

# Force UTF-8 output
sys.stdout = codecs.getwriter('utf-8')(sys.stdout.buffer, 'replace')

try:
    import docx
except ImportError:
    print("Installing python-docx...")
    os.system("pip install python-docx")
    import docx

# Paths
SOURCE_DIR = Path("C:/Users/edbar/Downloads/Character/Charactere")
OUTPUT_DIR = Path(__file__).parent / "character_exports"

def extract_document(filepath: Path) -> dict:
    """Extract all text content from a .docx file."""
    doc = docx.Document(str(filepath))

    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Detect if paragraph is bold (heading/NPC name)
            is_bold = False
            if para.runs:
                is_bold = all(run.bold for run in para.runs if run.text.strip())

            paragraphs.append({
                'text': text,
                'is_bold': is_bold,
                'style': para.style.name if para.style else None
            })

    # Also extract tables if any
    tables = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        if table_data:
            tables.append(table_data)

    return {
        'filename': filepath.name,
        'paragraphs': paragraphs,
        'tables': tables,
        'raw_text': '\n\n'.join(p['text'] for p in paragraphs)
    }

def main():
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    # Get all docx files
    docx_files = sorted(SOURCE_DIR.glob("*.docx"))
    print(f"Found {len(docx_files)} documents\n")

    results = []

    for filepath in docx_files:
        char_name = filepath.stem
        print(f"Extracting: {char_name}")

        try:
            data = extract_document(filepath)

            # Save raw text
            txt_path = OUTPUT_DIR / f"{char_name}.txt"
            with open(txt_path, 'w', encoding='utf-8') as f:
                f.write(data['raw_text'])

            # Save structured data
            json_path = OUTPUT_DIR / f"{char_name}.json"
            with open(json_path, 'w', encoding='utf-8') as f:
                json.dump(data, f, indent=2, ensure_ascii=False)

            results.append({
                'name': char_name,
                'paragraphs': len(data['paragraphs']),
                'chars': len(data['raw_text']),
                'success': True
            })

            print(f"  Paragraphs: {len(data['paragraphs'])}, Chars: {len(data['raw_text'])}")

        except Exception as e:
            print(f"  ERROR: {e}")
            results.append({
                'name': char_name,
                'error': str(e),
                'success': False
            })

    # Save summary
    summary_path = OUTPUT_DIR / "_extraction_summary.json"
    with open(summary_path, 'w', encoding='utf-8') as f:
        json.dump(results, f, indent=2)

    print(f"\n{'='*60}")
    print(f"Extraction complete!")
    print(f"Output directory: {OUTPUT_DIR}")
    print(f"Total documents: {len(results)}")
    print(f"Successful: {sum(1 for r in results if r.get('success'))}")

if __name__ == "__main__":
    main()
